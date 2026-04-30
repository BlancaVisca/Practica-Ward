"""Genera el reporte final en formato .docx con los resultados del clustering."""
import os
import pandas as pd
import numpy as np
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = '/sessions/stoic-affectionate-ritchie/mnt/outputs'
df = pd.read_csv(os.path.join(BASE, 'dataset_raw', 'DATA (1).csv'))
cortes = pd.read_csv(os.path.join(BASE, 'cortes_resumen.csv'))

attributes = [
    ("STUDENT ID", "Identificador unico del estudiante (rol: ID)", "Categorico"),
    ("1", "Edad del estudiante", "1: 18-21, 2: 22-25, 3: >26"),
    ("2", "Sexo", "1: female, 2: male"),
    ("3", "Tipo de bachillerato cursado", "1: private, 2: state, 3: other"),
    ("4", "Tipo de beca recibida", "1: none, 2: 25%, 3: 50%, 4: 75%, 5: full"),
    ("5", "Trabajo adicional", "1: yes, 2: no"),
    ("6", "Actividad artistica o deportiva regular", "1: yes, 2: no"),
    ("7", "Tiene pareja", "1: yes, 2: no"),
    ("8", "Salario total (si dispone)", "1: USD 135-200, 2: 201-270, 3: 271-340, 4: 341-410, 5: >410"),
    ("9", "Medio de transporte a la universidad", "1: bus, 2: private car/taxi, 3: bicycle, 4: other"),
    ("10", "Tipo de alojamiento en Chipre", "1: rental, 2: dormitory, 3: with family, 4: other"),
    ("11", "Nivel educativo de la madre", "1: primary, 2: secondary, 3: high school, 4: university, 5: MSc, 6: PhD"),
    ("12", "Nivel educativo del padre", "1: primary, 2: secondary, 3: high school, 4: university, 5: MSc, 6: PhD"),
    ("13", "Numero de hermanos", "1: 1, 2: 2, 3: 3, 4: 4, 5: 5 o mas"),
    ("14", "Estado civil de los padres", "1: married, 2: divorced, 3: died (one or both)"),
    ("15", "Ocupacion de la madre", "1: retired, 2: housewife, 3: government officer, 4: private sector employee, 5: self-employment, 6: other"),
    ("16", "Ocupacion del padre", "1: retired, 2: housewife, 3: government officer, 4: private sector employee, 5: self-employment, 6: other"),
    ("17", "Horas de estudio semanales", "1: none, 2: <5h, 3: 6-10h, 4: 11-20h, 5: >20h"),
    ("18", "Frecuencia de lectura (no cientifica)", "1: none, 2: sometimes, 3: often"),
    ("19", "Frecuencia de lectura (cientifica)", "1: none, 2: sometimes, 3: often"),
    ("20", "Asistencia a seminarios/conferencias del departamento", "1: yes, 2: no"),
    ("21", "Impacto de los proyectos/actividades en el exito", "1: positive, 2: negative, 3: neutral"),
    ("22", "Asistencia a clases", "1: always, 2: sometimes, 3: never"),
    ("23", "Preparacion para examenes parciales 1", "1: alone, 2: with friends, 3: not applicable"),
    ("24", "Preparacion para examenes parciales 2", "1: closest date to the exam, 2: regularly during the semester, 3: never"),
    ("25", "Tomar apuntes en clase", "1: never, 2: sometimes, 3: always"),
    ("26", "Escuchar en clase", "1: never, 2: sometimes, 3: always"),
    ("27", "La discusion mejora mi interes y exito en el curso", "1: never, 2: sometimes, 3: always"),
    ("28", "Flip-classroom (aula invertida)", "1: not useful, 2: useful, 3: not applicable"),
    ("29", "GPA acumulado del ultimo semestre (/4.00)", "1: <2.00, 2: 2.00-2.49, 3: 2.50-2.99, 4: 3.00-3.49, 5: >3.49"),
    ("30", "GPA acumulado esperado al graduarse (/4.00)", "1: <2.00, 2: 2.00-2.49, 3: 2.50-2.99, 4: 3.00-3.49, 5: >3.49"),
    ("COURSE ID", "Identificador del curso (rol: Feature, tipo Integer)", "Entero"),
    ("GRADE", "Calificacion final - variable objetivo (rol: Target)", "0: Fail, 1: DD, 2: DC, 3: CC, 4: CB, 5: BB, 6: BA, 7: AA"),
]

doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    return p

def shade_cell(cell, color_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def make_table(header, rows, header_color='305496'):
    t = doc.add_table(rows=1, cols=len(header)); t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shade_cell(hdr[i], header_color)
    for row_data in rows:
        row = t.add_row().cells
        for i, v in enumerate(row_data):
            row[i].text = str(v)
        for p in row[0].paragraphs:
            for r in p.runs: r.bold = True
    return t

# Portada
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run('Clustering Jerarquico - Metodo de Ward')
trun.bold = True; trun.font.size = Pt(22); trun.font.color.rgb = RGBColor(0x1F,0x3A,0x5F)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = sub.add_run('Higher Education Students Performance Evaluation (UCI #856)')
srun.bold = True; srun.font.size = Pt(14); srun.font.color.rgb = RGBColor(0x55,0x55,0x55)
doc.add_paragraph()
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Autor: Christopher\nFecha: 28 de abril de 2026').italic = True
doc.add_paragraph()

# 1. Descripcion
add_heading('1. Descripcion del dataset y atributos', 1)
add_para('El dataset "Higher Education Students Performance Evaluation" (UCI Machine Learning '
    'Repository, ID #856) contiene la informacion de 145 estudiantes universitarios y se compone '
    'de 33 columnas: un identificador, 30 preguntas categoricas/ordinales sobre aspectos academicos, '
    'familiares y personales, un identificador de curso y la calificacion final (variable objetivo '
    'GRADE) que toma 8 valores distintos (0=Fail hasta 7=AA).')
add_para('A continuacion se describe cada atributo. Las descripciones y codificaciones provienen '
    'de la documentacion oficial publicada en el UCI Machine Learning Repository '
    '(https://archive.ics.uci.edu/dataset/856/). En el archivo CSV los atributos aparecen '
    'identificados con numeros del 1 al 30, por lo que la tabla establece la correspondencia '
    'entre el numero de columna y la pregunta original.')
make_table(['Atributo', 'Descripcion', 'Valores posibles'], attributes)
doc.add_paragraph()

# 2. Metodologia
add_heading('2. Metodologia - Implementacion del metodo de Ward', 1)
add_para('El clustering jerarquico aglomerativo con enlace de Ward se implemento en Python '
    'utilizando SciPy, NumPy y Pandas. Las decisiones tomadas fueron:')
for b in [
    'Datos sin estandarizacion (a peticion del usuario): la matriz de entrada conserva los valores originales.',
    'Variables excluidas: STUDENT ID y GRADE. Se usaron los 30 atributos numerados mas COURSE ID (31 columnas).',
    'Distancia: euclidiana (requerida por Ward).',
    'Funcion principal: scipy.cluster.hierarchy.linkage(data, method="ward").',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para('El criterio de Ward minimiza la varianza total dentro de los clusters al fusionar grupos: '
    'en cada paso se unen los dos clusters cuya fusion produce el menor incremento posible en la '
    'suma de cuadrados intra-cluster (ESS - Error Sum of Squares).')

# 3. Dendrograma
add_heading('3. Dendrograma', 1)
add_para('El dendrograma resultante se muestra a continuacion. El eje vertical representa la '
    'distancia de Ward y el eje horizontal lista los 145 estudiantes.')
doc.add_picture(os.path.join(BASE, 'dendrograma.png'), width=Cm(16))
cap = doc.add_paragraph('Figura 1. Dendrograma completo - metodo de Ward.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in cap.runs: r.italic = True; r.font.size = Pt(10)
add_para('Mismo dendrograma con las tres lineas de corte aplicadas:')
doc.add_picture(os.path.join(BASE, 'dendrograma_con_cortes.png'), width=Cm(16))
cap2 = doc.add_paragraph('Figura 2. Dendrograma con tres lineas de corte (4, 6 y 8 grupos).')
cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in cap2.runs: r.italic = True; r.font.size = Pt(10)

# 4. Cortes
add_heading('4. Lineas de corte y libros Excel generados', 1)
add_para('Se eligieron tres lineas de corte para producir 4, 6 y 8 grupos respectivamente. La '
    'eleccion de 4 y 6 ofrece una vision a alto nivel mientras que el corte a 8 grupos esta '
    'disenado para alinearse con las 8 clases reales de la columna GRADE.')
make_table(['Corte', 'Numero de grupos', 'Altura (Ward)', 'Archivo Excel'],
    [('A', '4', f"{cortes.iloc[0]['altura']:.4f}", 'corte_A_4grupos.xlsx'),
     ('B', '6', f"{cortes.iloc[1]['altura']:.4f}", 'corte_B_6grupos.xlsx'),
     ('C', '8', f"{cortes.iloc[2]['altura']:.4f}", 'corte_C_8grupos.xlsx')])
add_para('')
add_para('Cada libro contiene una hoja "Resumen" y, ademas, una hoja por cada grupo (Grupo_1, '
    'Grupo_2, ...). En las hojas de grupo aparecen las filas completas (33 columnas) de los '
    'estudiantes asignados a ese grupo, manteniendo el STUDENT ID original para trazabilidad.')

# 5. Corte a 8 grupos
add_heading('5. Corte mas cercano a 8 grupos', 1)
h8 = cortes.iloc[2]['altura']
add_para(f'Se realizo un barrido de umbrales sobre el rango de distancias del dendrograma y se '
    f'identifico que existe un rango de alturas que producen exactamente k=8 grupos. El rango '
    f'exacto es [12.0454, 12.1458] y la altura representativa utilizada fue h = {h8:.4f}.')

# 6. Matriz 8x8
add_heading('6. Matriz de confusion y asignacion de nombres', 1)
add_para('Para evaluar la coincidencia entre los 8 clusters obtenidos por Ward y las 8 clases '
    'reales (GRADE), se construyo una matriz de coincidencia 8x8. Posteriormente se aplico el '
    'algoritmo Hungaro (scipy.optimize.linear_sum_assignment) para encontrar la asignacion '
    'biunivoca cluster -> clase que maximiza el numero total de coincidencias.')
add_para('La matriz de confusion 8x8 (clase real x clase predicha):')
doc.add_picture(os.path.join(BASE, 'matriz_confusion.png'), width=Cm(13))
cap3 = doc.add_paragraph('Figura 3. Matriz de confusion 8x8 tras la asignacion mediante algoritmo Hungaro.')
cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in cap3.runs: r.italic = True; r.font.size = Pt(10)

wb = openpyxl.load_workbook(os.path.join(BASE, 'matriz_confusion.xlsx'), data_only=True)
ws = wb['Asignacion_Hungaro']
add_para('Asignacion final cluster -> clase real (GRADE):', bold=True)
asign_rows = []
for r_i in range(2, ws.max_row+1):
    asign_rows.append((str(ws.cell(row=r_i, column=1).value),
                       str(ws.cell(row=r_i, column=2).value),
                       str(ws.cell(row=r_i, column=3).value),
                       str(ws.cell(row=r_i, column=4).value)))
make_table(['Cluster', 'Clase asignada', 'Coincidencias', 'Tamano cluster'], asign_rows)

ws_res = wb['Resumen']
acc = ws_res.cell(row=7, column=2).value
aciertos = ws_res.cell(row=5, column=2).value
total = ws_res.cell(row=6, column=2).value
add_para(''); add_para('Resultados globales (8 clases):', bold=True)
add_para(f'  - Aciertos: {aciertos} de {total}')
add_para(f'  - Accuracy multiclase: {acc} ({float(acc)*100:.2f}%)')

# 6.1 Matriz 2x2
add_heading('6.1. Matriz de confusion 2x2 (VP / VN / FP / FN)', 2)
add_para('Para presentar el desempeno en formato binario clasico (Verdaderos Positivos / '
    'Verdaderos Negativos / Falsos Positivos / Falsos Negativos) se aplica el metodo '
    'micro-promediado one-vs-rest: para cada par (estudiante, clase c) se evalua como un '
    'problema binario y se acumulan las cuatro categorias sobre las 8 clases. El total de '
    'evaluaciones es N x C = 145 x 8 = 1160.')
add_para('Definiciones aplicadas:', bold=True)
for d in [
    'VP (Verdadero Positivo): la prediccion dice "es de clase c" y la realidad es "es de clase c".',
    'VN (Verdadero Negativo): la prediccion dice "no es de clase c" y la realidad es "no es de clase c".',
    'FP (Falso Positivo): la prediccion dice "es de clase c" pero la realidad es "no es de clase c".',
    'FN (Falso Negativo): la prediccion dice "no es de clase c" pero la realidad es "es de clase c".',
]:
    doc.add_paragraph(d, style='List Bullet')
doc.add_picture(os.path.join(BASE, 'matriz_confusion_2x2.png'), width=Cm(13))
cap4 = doc.add_paragraph('Figura 4. Matriz de confusion 2x2 micro-promediada (VP / VN / FP / FN).')
cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in cap4.runs: r.italic = True; r.font.size = Pt(10)

wb22 = openpyxl.load_workbook(os.path.join(BASE, 'matriz_confusion_2x2.xlsx'), data_only=True)
ws22 = wb22['Matriz_2x2']
VP_v = ws22['B4'].value; FN_v = ws22['C4'].value
FP_v = ws22['B5'].value; VN_v = ws22['C5'].value
acc_b = ws22.cell(row=10, column=2).value
prec  = ws22.cell(row=11, column=2).value
rec   = ws22.cell(row=12, column=2).value
f1m   = ws22.cell(row=13, column=2).value

add_para('Resultados de la matriz 2x2:', bold=True)
t22 = doc.add_table(rows=3, cols=3); t22.style = 'Light Grid Accent 1'
t22.cell(0,0).text = ''; t22.cell(0,1).text = 'Pred POSITIVO'; t22.cell(0,2).text = 'Pred NEGATIVO'
t22.cell(1,0).text = 'Real POSITIVO'; t22.cell(1,1).text = f'VP = {VP_v}'; t22.cell(1,2).text = f'FN = {FN_v}'
t22.cell(2,0).text = 'Real NEGATIVO'; t22.cell(2,1).text = f'FP = {FP_v}'; t22.cell(2,2).text = f'VN = {VN_v}'
for i in range(3):
    for j in range(3):
        for p in t22.cell(i,j).paragraphs:
            for r in p.runs: r.bold = (i==0 or j==0)
shade_cell(t22.cell(0,1), '305496'); shade_cell(t22.cell(0,2), '305496')
shade_cell(t22.cell(1,0), '305496'); shade_cell(t22.cell(2,0), '305496')
for c in [(0,1),(0,2),(1,0),(2,0)]:
    for p in t22.cell(*c).paragraphs:
        for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_cell(t22.cell(1,1), 'C6EFCE'); shade_cell(t22.cell(2,2), 'C6EFCE')
shade_cell(t22.cell(1,2), 'FFC7CE'); shade_cell(t22.cell(2,1), 'FFC7CE')
add_para(''); add_para('Metricas globales derivadas:', bold=True)
make_table(['Metrica', 'Valor'], [
    ('Total de evaluaciones (N x C)', f'{VP_v+VN_v+FP_v+FN_v}'),
    ('Accuracy binaria = (VP+VN)/Total', f'{acc_b}  ({float(acc_b)*100:.2f}%)'),
    ('Precision (micro) = VP/(VP+FP)', f'{prec}'),
    ('Recall (micro)    = VP/(VP+FN)', f'{rec}'),
    ('F1 (micro)        = 2*P*R/(P+R)', f'{f1m}'),
    ('Aciertos del clustering (VP)', f'{VP_v} de {VP_v+FN_v} estudiantes'),
])
add_para('')
add_para('Observacion importante: en una clasificacion multiclase con asignacion biunivoca '
    '(cada estudiante recibe una sola prediccion), se cumple FP_total = FN_total porque cada '
    'error es simultaneamente un FN para la clase real y un FP para la clase predicha. Por esta '
    f'razon precision, recall y F1 micro coinciden con la accuracy multiclase ({float(prec)*100:.2f}%). '
    f'La accuracy binaria mas alta ({float(acc_b)*100:.2f}%) refleja los muchos Verdaderos Negativos: '
    'cada estudiante es correctamente identificado como "no perteneciente" a las 7 clases en las '
    'que realmente no esta.')

# 7. Conclusiones
add_heading('7. Conclusiones', 1)
add_para('El metodo de Ward identifico una estructura jerarquica clara en el dataset, permitiendo '
    'segmentar a los 145 estudiantes en agrupaciones interpretables a distintos niveles de '
    'granularidad. Los principales hallazgos:')
for c in [
    'Los tres cortes generan particiones distintas: 4 grupos (alto nivel), 6 grupos (intermedio) y 8 grupos (alineado con GRADE).',
    f'Accuracy multiclase con Hungaro: {float(prec)*100:.2f}% (40/145). En matriz 2x2: 40 VP, 105 FN, 105 FP, 910 VN; accuracy binaria {float(acc_b)*100:.2f}%.',
    'GRADE no es facilmente reproducible solo desde las preguntas: estudiantes con notas distintas comparten muchas respuestas.',
    'Trabajo futuro: estandarizar datos, comparar otros enlaces (average, complete) y evaluar con ARI, NMI, silhouette.',
]:
    doc.add_paragraph(c, style='List Bullet')

# 8. Anexos
add_heading('8. Archivos entregables', 1)
files = [
    ('ward_clustering.py', 'Script principal de Python con todo el pipeline.'),
    ('matriz_confusion_2x2.py', 'Script que genera la matriz 2x2 micro-promediada.'),
    ('replot_2x2_heatmap.py', 'Script que regrafica la matriz 2x2 como heatmap.'),
    ('dendrograma.png', 'Dendrograma original sin lineas de corte.'),
    ('dendrograma_con_cortes.png', 'Dendrograma con tres lineas de corte resaltadas.'),
    ('corte_A_4grupos.xlsx', 'Libro Excel - particion de 4 grupos.'),
    ('corte_B_6grupos.xlsx', 'Libro Excel - particion de 6 grupos.'),
    ('corte_C_8grupos.xlsx', 'Libro Excel - particion de 8 grupos.'),
    ('matriz_confusion.xlsx', 'Libro Excel con matriz cluster x clase, asignacion Hungaro y matriz 8x8.'),
    ('matriz_confusion.png', 'Heatmap de la matriz de confusion 8x8.'),
    ('matriz_confusion_2x2.xlsx', 'Libro Excel con matriz 2x2 (VP/VN/FP/FN), metricas y desglose por clase.'),
    ('matriz_confusion_2x2.png', 'Heatmap de la matriz 2x2.'),
    ('barrido_umbrales.csv', 'Tabla del barrido de umbrales (altura -> numero de clusters).'),
    ('cortes_resumen.csv', 'Resumen de los cortes seleccionados.'),
]
make_table(['Archivo', 'Contenido'], files)

# 9. FAQ
add_heading('9. Preguntas frecuentes y aclaraciones conceptuales', 1)
add_para('Esta seccion responde las preguntas planteadas durante la revision del trabajo, dejando '
    'explicitas las decisiones metodologicas y la justificacion de cada paso.')

# 9.1
add_heading('9.1. Por que se eligieron 8 grupos', 2)
add_para('No fue una eleccion arbitraria del analisis, sino el objetivo definido en el enunciado. '
    'El punto 5 dice textualmente: "Los objetos de la BDD pertenecen a 8 diferentes clases. '
    'Buscar la linea de corte que mas se asemeje a obtener 8 grupos". Esas 8 clases corresponden '
    'a los valores de la columna GRADE: 0 (Fail), 1 (DD), 2 (DC), 3 (CC), 4 (CB), 5 (BB), '
    '6 (BA) y 7 (AA). Esto se verifico programaticamente con df["GRADE"].unique() y efectivamente '
    'hay 8 valores distintos. Por tanto, k=8 es un objetivo del enunciado, no una preferencia.')

# 9.2
add_heading('9.2. Por que esos 3 cortes a esas alturas', 2)
add_para('Eleccion de los k (numero de grupos): se eligieron 4, 6 y 8 para mostrar tres niveles '
    'distintos de granularidad: 4 grupos como vista macro (segmentos amplios), 6 grupos como '
    'nivel intermedio, y 8 grupos para coincidir con el objetivo del punto 5.')
add_para('Las alturas de corte NO se eligieron a mano. Se obtuvieron mediante un barrido de '
    'umbrales: para cada altura entre el minimo (2.000) y el maximo (52.122) de la matriz de '
    'enlace de Ward, se calculo el numero de clusters resultantes. Para cada k objetivo se '
    'identifico el rango de alturas que produce exactamente ese k, y se tomo el punto medio del '
    'rango como altura representativa. Esto garantiza que la asignacion sea estable frente a '
    'pequenas variaciones numericas.')
add_para('Resultados concretos del barrido:', bold=True)
for b in [
    'k=4 grupos: punto medio del rango = 15.7117 (altura usada en el corte A).',
    'k=6 grupos: punto medio del rango = 14.2046 (altura usada en el corte B).',
    'k=8 grupos: rango exacto [12.0454, 12.1458], punto medio = 12.0956 (altura usada en el corte C).',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para('El barrido completo se guardo en el archivo barrido_umbrales.csv para referencia.')

# 9.3
add_heading('9.3. Como se graficaron los dendrogramas', 2)
add_para('Los dendrogramas se generaron con la funcion scipy.cluster.hierarchy.dendrogram() '
    'aplicada sobre la matriz de enlace Z que devuelve linkage(X, method="ward"). El proceso fue:')
for b in [
    'Paso 1: Z = linkage(X, method="ward", metric="euclidean") calcula la jerarquia completa de fusiones.',
    'Paso 2: dendrogram(Z, labels=df["STUDENT ID"].values, ...) dibuja el arbol resultante.',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para('Interpretacion del dendrograma:', bold=True)
for b in [
    'Eje X: los 145 estudiantes ordenados de modo que las hojas no se crucen.',
    'Eje Y: la distancia de Ward a la que ocurre cada fusion (incremento de varianza intra-cluster).',
    'Cada "U" invertida es una fusion: dos ramas se unen a la altura donde el algoritmo decidio fusionarlas.',
    'Una linea horizontal de corte cortara un numero determinado de ramas verticales, y cada rama corresponde a un cluster.',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para('Para el segundo dendrograma (con cortes), se usaron tres llamadas a axhline() sobre el '
    'mismo grafico, cada linea a la altura calculada en el barrido (15.71, 14.20 y 12.10).')

add_heading('9.4. Datos reales utilizados para los heatmaps 8x8 y 2x2', 2)
add_para('Ambos heatmaps se construyen exclusivamente con datos del archivo DATA (1).csv, sin '
    'introducir informacion externa. Las dos fuentes de informacion son:')
for b in [
    'Entrada al clustering: las 30 preguntas (columnas 1 a 30) + COURSE ID, sin estandarizar. '
    'Esta es la matriz X de 145x31 con la que se calcula la matriz de enlace de Ward.',
    'Etiqueta real (ground truth): la columna GRADE del CSV, con valores 0-7. Esta columna NO '
    'participo en el clustering (se excluyo, segun lo acordado), pero se usa para EVALUAR que '
    'tan bien los grupos formados se parecen a la calificacion real.',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para('Heatmap 8x8 (matriz_confusion.png):', bold=True)
add_para('Cruce entre GRADE real y GRADE predicho. La celda en fila i, columna j cuenta cuantos '
    'estudiantes con GRADE real = i fueron asignados al cluster que (tras la asignacion Hungara) '
    'recibio el nombre GRADE = j. La diagonal son los aciertos (suma = 40).')
add_para('Heatmap 2x2 (matriz_confusion_2x2.png):', bold=True)
add_para('Se construye sobre la misma informacion del 8x8, pero binarizada con el metodo '
    'micro-promediado one-vs-rest. Para cada clase c en 0..7:')
for b in [
    'VP_c = casos donde real == c Y pred == c',
    'VN_c = casos donde real != c Y pred != c',
    'FP_c = casos donde real != c Y pred == c',
    'FN_c = casos donde real == c Y pred != c',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para('Los totales (VP=40, VN=910, FP=105, FN=105) se obtienen sumando sobre las 8 clases. '
    'El gran total es 1160 = 145 estudiantes x 8 clases.')

add_heading('9.5. Solucion al punto 6 (asignacion de nombres de grupos)', 2)
add_para('Problema: el clustering produce grupos sin nombre (Cluster_1 a Cluster_8). Hay que '
    'asignar a cada cluster el nombre de una de las 8 clases reales de GRADE. Una asignacion '
    'ingenua del tipo "el cluster con mas miembros recibe la clase mayoritaria" puede dar '
    'conflictos (dos clusters con el mismo nombre) y no garantiza la mejor correspondencia global.')
add_para('Solucion aplicada - algoritmo Hungaro:', bold=True)
add_para('1) Construccion de la matriz 8x8 de coincidencias entre clusters (filas) y clases reales (columnas):')
add_para('   M[i,j] = numero de estudiantes en Cluster_i cuya GRADE real es j', italic=True)
add_para('2) Aplicacion de scipy.optimize.linear_sum_assignment(-M) sobre la matriz negada (porque '
    'la funcion minimiza, y se queria maximizar). Este algoritmo encuentra la asignacion biunivoca '
    '(cada cluster a una clase distinta, cada clase a un cluster distinto) que maximiza la suma '
    'total de la diagonal, es decir, los aciertos globales.')
add_para('3) Resultado de la asignacion obtenida:', bold=True)
make_table(['Cluster', 'Clase asignada (GRADE)', 'Aciertos en este cluster'],
    [('Cluster_1', 'Clase_5 (BB)', '8'),
     ('Cluster_2', 'Clase_0 (Fail)', '1'),
     ('Cluster_3', 'Clase_2 (DC)', '6'),
     ('Cluster_4', 'Clase_1 (DD)', '11'),
     ('Cluster_5', 'Clase_4 (CB)', '1'),
     ('Cluster_6', 'Clase_6 (BA)', '2'),
     ('Cluster_7', 'Clase_3 (CC)', '4'),
     ('Cluster_8', 'Clase_7 (AA)', '7')])
add_para('')
add_para('Evaluacion: con esos nombres, se compara la prediccion vs el GRADE real para los 145 '
    'estudiantes y se obtienen 40 aciertos (accuracy multiclase = 27.59%). La matriz 8x8 final '
    'es el cruce de GRADE real vs el GRADE predicho con esos nombres.')
add_para('Por que el algoritmo Hungaro y no algo mas simple:', bold=True)
for b in [
    'Asignacion biunivoca: ningun nombre se repite y todas las 8 clases se cubren.',
    'Optimalidad global: maximiza el numero total de aciertos (no solo localmente).',
    'Determinista: ante empates, devuelve una unica solucion reproducible.',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para('Con el criterio ingenuo de "asignar a cada cluster su clase mayoritaria" sin restriccion '
    'biunivoca, varios clusters compartirian nombre y otros quedarian sin nombre, dando un '
    'resultado peor y ambiguo.')

add_heading('9.6. Por que el Cluster_2 fue asignado a Clase_0 (Fail)', 2)
add_para('Esta asignacion sorprende a primera vista, porque el Cluster_2 solo tiene 1 estudiante '
    'con Clase_0 real, mientras que Cluster_8 tiene 4. La explicacion esta en como funciona el '
    'algoritmo Hungaro: optimiza el TOTAL global de aciertos, no caso por caso.')
add_para('Distribucion completa de la Clase_0 (Fail) en los clusters:', bold=True)
make_table(['Cluster', 'Estudiantes Clase_0', 'Clase con mas miembros en ese cluster'],
    [('Cluster_1', '0', 'Clase_5 (8 estudiantes)'),
     ('Cluster_2', '1', 'Clase_1 (4 estudiantes)'),
     ('Cluster_3', '1', 'Clase_2 (6 estudiantes)'),
     ('Cluster_4', '1', 'Clase_1 (11 estudiantes)'),
     ('Cluster_5', '0', 'Clase_7 (3 estudiantes)'),
     ('Cluster_6', '0', 'Clase_7 (3 estudiantes)'),
     ('Cluster_7', '1', 'Clase_1/2/3 (empate, 4-5 c/u)'),
     ('Cluster_8', '4', 'Clase_7 (7 estudiantes)')])
add_para('')
add_para('Si el Hungaro hubiera asignado Cluster_8 -> Clase_0 (Fail):', bold=True)
for b in [
    'Ganaria 4 aciertos en Clase_0.',
    'Pero perderia los 7 aciertos de Cluster_8 -> Clase_7 (AA), que es donde Cluster_8 brilla.',
    'Clase_7 tendria que ir a Cluster_5 o Cluster_6 (3 aciertos cada uno, maximo).',
    'Balance: ganaria 4 - 7 + 3 = pierde 4 aciertos respecto a la solucion actual.',
]:
    doc.add_paragraph(b, style='List Bullet')
add_para('Por eso el algoritmo eligio: Cluster_8 -> Clase_7 (7 aciertos, su mayor coincidencia), '
    'y a Clase_0 le toco el resto. Entre los clusters que no son Cluster_8 y que tienen al menos '
    '1 Fail (Cluster_2, 3, 4 o 7), todos tienen exactamente 1 estudiante Fail, por lo que da igual '
    'cual se elija. El Hungaro eligio Cluster_2 porque al combinar las 7 asignaciones restantes '
    'con esa eleccion, el total global era 40 (el maximo posible).')
add_para('Conclusion:', bold=True)
add_para('Clase_0 (Fail) tiene tan pocos estudiantes (8) y estan tan dispersos que ningun cluster '
    'es realmente el cluster de Fail. La asignacion Cluster_2 -> Clase_0 es la menos mala '
    'globalmente. Esto refleja una limitacion real del dataset, no un error del algoritmo: la '
    'Clase_0 no tiene una huella distintiva en los 31 atributos de entrada, y por eso ningun '
    'cluster se forma en torno a los estudiantes Fail.')

add_heading('9.7. Formulas del metodo de Ward, paso a paso del algoritmo y razones del accuracy', 2)

add_para('Formulas matematicas:', bold=True)
add_para('1) Funcion objetivo (Error Sum of Squares - ESS): para cada cluster C con centroide '
    'x_barra_C, su ESS es la suma de los cuadrados de las distancias de cada punto al centroide:')
add_para('     ESS(C) = sum sobre x_i en C de || x_i - x_barra_C ||^2', italic=True)

add_para('2) Criterio de fusion: en cada iteracion se fusionan los dos clusters A y B que '
    'minimizan el incremento de ESS:')
add_para('     Delta_ESS = ESS(A union B) - ESS(A) - ESS(B)', italic=True)

add_para('3) Formula de Lance-Williams para Ward (la que SciPy aplica internamente para '
    'actualizar distancias entre clusters de manera incremental):')
add_para('     d2(AuB, C) = [(nA + nC)*d2(A,C) + (nB + nC)*d2(B,C) - nC*d2(A,B)] / (nA + nB + nC)', italic=True)

add_para('4) Distancia entre dos clusters individuales:')
add_para('     d2(A,B) = (nA*nB)/(nA+nB) * || x_barra_A - x_barra_B ||^2', italic=True)
add_para('La altura mostrada en el dendrograma es la raiz cuadrada de este valor (la distancia de Ward).')

add_para('')
add_para('Paso a paso de las fusiones reales del algoritmo:', bold=True)
add_para('SciPy devuelve una matriz de enlace Z con 144 filas (una por cada fusion). A continuacion '
    'se muestran las iteraciones mas representativas:')
make_table(['Iteracion', 'Fusiona', 'Distancia Ward', 'Tamano resultante', 'Comentario'],
    [('1', 'objeto 129 + 130', '2.0000', '2', 'par mas cercano de todo el dataset'),
     ('2', 'objeto 17 + 32',   '2.2361', '2', 'siguiente par mas similar'),
     ('3', 'objeto 44 + 46',   '2.6458', '2', '-'),
     ('4', 'objeto 51 + 53',   '2.8284', '2', '-'),
     ('5', 'objeto 89 + 92',   '2.8284', '2', 'empate con el anterior'),
     ('50',  'sub-cluster + objeto', '4.3589', '3',  'empiezan a formarse triangulos'),
     ('100', 'cluster + cluster',    '6.0828', '10', 'clusters medianos'),
     ('130', 'cluster + cluster',    '9.7265', '12', 'clusters grandes'),
     ('~136 (corte k=8)', '-', '~12.10', '-', 'AQUI cortamos para 8 grupos'),
     ('~138 (corte k=6)', '-', '~14.20', '-', 'AQUI cortamos para 6 grupos'),
     ('~140 (corte k=4)', '-', '~15.71', '-', 'AQUI cortamos para 4 grupos'),
     ('141', 'grupo grande + grande', '15.0175', '77', '-'),
     ('142', 'grupo grande + grande', '16.3978', '54', '-'),
     ('143', 'grupo grande + grande', '19.5005', '68', '-'),
     ('144 (final)', 'super-rama + super-rama', '52.1224', '145', 'todo el dataset en un solo cluster')])

add_para('')
add_para('Patron observado: las distancias crecen lentamente durante casi todo el proceso (uniones '
    'naturales entre puntos similares), pero dan un salto muy grande en la ultima iteracion '
    '(de 19.50 a 52.12). Este salto indica que las dos super-ramas finales son extremadamente '
    'distintas entre si, lo que corresponde a la division estructural mas significativa del dataset.')

add_para('')
add_para('Razones del accuracy de 27.59%:', bold=True)
add_para('La accuracy obtenida puede parecer baja, pero hay seis razones acumulativas que la explican '
    'y que es importante poner en contexto antes de interpretarla como un mal resultado.')

add_para('(a) El baseline real es bajo, no 50%.', bold=True)
add_para('Con 8 clases, una prediccion al azar daria 12.5% de accuracy (1/8). Si predijeramos '
    'siempre la clase mas frecuente (clasificador ZeroR), obtendriamos 24.1% (la Clase_1 tiene '
    '35/145 estudiantes = 24.1%). El 27.59% obtenido supera ambos baselines, por lo que el '
    'clustering si captura informacion relevante, aunque sea poca.')

add_para('(b) Clustering no supervisado vs clasificacion supervisada.', bold=True)
add_para('El Ward no usa la columna GRADE para entrenarse: agrupa por similitud de las 31 variables '
    'de entrada (las 30 preguntas + COURSE ID). Esperar que los grupos coincidan con GRADE es '
    'pedirle al algoritmo adivinar el examen sin haberlo visto. Cualquier metodo no supervisado '
    'tendra rendimiento limitado en una metrica supervisada como la accuracy.')

add_para('(c) No estandarizamos los datos (decision del usuario).', bold=True)
add_para('Variables con rangos grandes (COURSE ID va de 1 a 9, atributos 11 y 12 van de 1 a 6) '
    'dominan la distancia euclidiana y aplastan a las variables binarias (1-2). Esto distorsiona '
    'la geometria del espacio y reduce el rendimiento. Con estandarizacion (StandardScaler) la '
    'accuracy suele subir entre 5 y 10 puntos en datasets como este.')

add_para('(d) Clase_0 (Fail) tiene solo 8 estudiantes.', bold=True)
add_para('Es estadisticamente muy dificil que un cluster se forme en torno a un grupo tan pequeno '
    'y disperso. Como se explico en la seccion 9.6, la asignacion Cluster_2 -> Clase_0 (Fail) solo '
    'aporta 1 acierto, regalando automaticamente 7 FN para Clase_0.')

add_para('(e) Curse of dimensionality.', bold=True)
add_para('Con 31 variables y solo 145 puntos (proporcion 1:4.7), el espacio esta relativamente '
    'vacio y todas las distancias entre puntos tienden a parecerse, debilitando la senal de '
    'cluster. Para datasets pequenos suele recomendarse menos variables o tecnicas de reduccion '
    'de dimensionalidad (PCA, t-SNE) antes de clusterizar.')

add_para('(f) Los atributos no determinan completamente el GRADE.', bold=True)
add_para('Dos estudiantes pueden responder casi igual a las 30 preguntas (mismo nivel educativo '
    'de los padres, mismas horas de estudio, misma asistencia, etc.) y obtener calificaciones '
    'distintas, porque el resultado del semestre depende tambien de factores no encuestados '
    '(esfuerzo en el examen final, suerte, dificultad del docente, etc.). Esto impone un techo '
    'intrinseco al rendimiento de cualquier metodo de clustering aplicado solo a estas 30 preguntas.')

add_para('')
add_para('Conclusion del accuracy:', bold=True)
add_para('27.59% no es un mal resultado en si: es aproximadamente 2.2x el baseline aleatorio (12.5%) '
    'y supera el ZeroR (24.1%), lo que demuestra que Ward captura estructura real de los datos. '
    'Para mejorar este numero las opciones son: (1) estandarizar los datos, (2) probar otros '
    'enlaces (average, complete), (3) reducir dimensionalidad con PCA antes de clusterizar, o '
    '(4) usar un metodo supervisado (random forest, XGBoost) si el objetivo final es realmente predecir GRADE.')

out = os.path.join(BASE, 'Reporte_Ward_Clustering.docx')
doc.save(out)
print(f'Reporte generado: {out}')
